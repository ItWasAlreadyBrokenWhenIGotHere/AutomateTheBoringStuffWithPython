#!/usr/bin/env python3
"""
Module Docstring
"""

__author__ = "Petri Weckström"
__version__ = "0.1.0"
__license__ = "MIT"

import docx
import logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.disable(logging.CRITICAL) # This disable all logs as CRITICAL level is highest


def main():
    """ Main entry point of the app """
    d = docx.Document()  # returns document object
    d.add_paragraph('Hello this is a paragraph.')
    d.add_paragraph('This is another paragraph.')
    p = d.paragraphs[0]
    p.add_run('This is a new run.')
    p.runs[1].bold = True

    d.save('E:\\Git\\automation\\document_handling\\documents\\new_demo.docx')


if __name__ == "__main__":
    """ This is executed when run from the command line """
    main()
