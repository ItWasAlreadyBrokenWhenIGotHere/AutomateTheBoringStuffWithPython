#!/usr/bin/env python3
"""
Module Docstring for example how to read and edit Word document
"""

__author__ = "Petri Weckström"
__version__ = "0.1.0"
__license__ = "MIT"

import docx
import logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.disable(logging.CRITICAL) # This disable all logs as CRITICAL level is highest


def main():
    """ Main entry point of the app """
    d = docx.Document('E:\\Git\\automation\\document_handling\\documents\\demo.docx')  # returns document object
    print(d.paragraphs[0].text)
    print(d.paragraphs[1].text)

    p = d.paragraphs[1]
    print(p.runs[0].text)
    print(p.runs[1].text)
    print(p.runs[2].text)
    print(p.runs[3].text)

    # modify text; add underline and change text content
    p.runs[3].underline = True
    p.runs[3].text = 'italic and underlined'
    d.save('E:\\Git\\automation\\document_handling\\documents\\demo2.docx')




if __name__ == "__main__":
    """ This is executed when run from the command line """
    main()
